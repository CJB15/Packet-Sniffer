import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from mpl_toolkits.basemap import Basemap
from docx import Document
from pathlib import Path
import ipinfo
import sys

def graphgen(log_df, GUI, loadfile, ipadress, maxpackets): # This whole file is a mess, very sorry
	
	ip_src_df = None # Set Later, Holds all the packets where the specified ip adress was the sender
	ip_dst_df = None # Set Later, Holds all the packets where the specified ip adress was the reciever

	if maxpackets == None: # If the user did not set maxpackets defualt it to 50
		maxpackets = 50
	
	if(ipadress == None): # If no IP adresss has been specifed then generate outputs for all the data
		print("Generating Output Data")
		savefolder = "./" + loadfile + " Output"
	else: # If an IP adresss has been specifed then generate outputs only for packets sent to or from that IP adress
		print("Generating Output Data for IP " + ipadress)
		ipnickname = ipadress.split(':')
		savefolder = "./" + loadfile + " Output " + ipnickname[0]
		
		ip_src_df = log_df[log_df['src_ip'] == ipadress] # Get all the packets where the specified ip adress was the sender
		ip_dst_df = log_df[log_df['dst_ip'] == ipadress] # Get all the packets where the specified ip adress was the reciever
		
		if(ip_src_df.empty and ip_dst_df.empty): # If specified IP address never appears then flag error and end program
			print("Error: Specifed IP address doesn't occur.")
			sys.exit(1)

	loadfile = os.path.splitext(loadfile)[0] # remove the .pcap from the users entered file
	Path(savefolder).mkdir(parents=True, exist_ok=True) # MAke folder to save the outputs in
	
	sourceIPChart(log_df, ipadress, maxpackets, ip_src_df, ip_dst_df, savefolder) # Generate source IP chart
	
	destinationIPChart(log_df, ipadress, maxpackets, ip_src_df, ip_dst_df, savefolder) # Generate destination IP chart
	
	sourcePortChart(log_df, ipadress, maxpackets, ip_src_df, ip_dst_df, savefolder) # Generate source port chart
	
	destinationPortChart(log_df, ipadress, maxpackets, ip_src_df, ip_dst_df, savefolder) # Generate destination port chart
	
	ipinfo_token = "bd8a79d2129512" # Token to use IPinfo, !!! Note: if another user plans to use this code replace token with your own, this will work but is reliant of cjb15 account. !!!
	handler = ipinfo.getHandler(ipinfo_token)
	ip_data = {}
	try: # Uses IPinfo to look up the details on very ip adress
		for x in pd.unique(log_df[['src_ip', 'dst_ip']].values.ravel()):
			if x != "Unknown IP(s)":
				ipdetails = handler.getDetails(x)
				ip_data[x] = ipdetails
		print("IPinfo connection successful")
		ipinfo_works = True # Notes if successful
	except:
		print("IPinfo connection unsuccessful")
		ipinfo_works = False # Notes if not
	
	if(ipinfo_works): # If successful then generate some things it requires that info for
		sourceIPMap(log_df, ipadress, maxpackets, ip_src_df, ip_dst_df, ip_data, savefolder) # Generate map if packaet sorce locations # TODO combine these 2 to reduce code duplication
		destinationIPMap(log_df, ipadress, maxpackets, ip_src_df, ip_dst_df, ip_data, savefolder) # Generate map if packaet destination locations 
	else: # If unsuccsesful then skip these
		print("Can't generate IP location maps")
	
	if(ipadress == None):
		documentGeneric(log_df, ipadress, maxpackets, ip_src_df, ip_dst_df, ip_data, savefolder, loadfile, ipinfo_works) # Generate generic document on all the data from the pcap file # TODO combine these 2 to reduce code duplication
	
	else:
		documentSpecific(log_df, ipadress, maxpackets, ip_src_df, ip_dst_df, ip_data, savefolder, loadfile, ipinfo_works) # Generate specific document on the data relating to one IP address from the pcap file
		
	
	print('Outputs saved in the "' + savefolder + '" folder')
	

def sourceIPChart(log_df, ipadress, maxpackets, ip_src_df, ip_dst_df, savefolder):
	plt.figure(figsize=(18,12)) # Set size of all the graphs/maps
	if(ipadress == None): # Generate a bar chart showing the number of packets that were sent from different ip adresses 
		log_df['src_ip'].value_counts().nlargest(maxpackets).plot(kind = 'barh')
		plt.title("Source IP adresses of packets")
	else: # Generate a bar chart showing the number of packets that were sent from different ip adresses to the specifed ip adress
		ip_dst_df['src_ip'].value_counts().nlargest(maxpackets).plot(kind = 'barh')
		plt.title("Source IP adresses of packets sent to " + ipadress)
	plt.xlabel("No. of Packets")
	plt.ylabel("IP Adress")
	plt.savefig(savefolder + "/src_ip.png", bbox_inches='tight')
	plt.clf()
	print("Source IP Graph Generated")

def destinationIPChart(log_df, ipadress, maxpackets, ip_src_df, ip_dst_df, savefolder):
	plt.figure(figsize=(18,12)) # Set size of all the graphs/maps
	if(ipadress == None): # Generate a bar chart showing the number of packets that were sent to different ip adresses
		log_df['dst_ip'].value_counts().nlargest(maxpackets).plot(kind = 'barh')
		plt.title("Destination IP adresses of packets")
	else: # Generate a bar chart showing the number of packets that were sent to different ip adresses from the specifed ip adress
		ip_src_df['dst_ip'].value_counts().nlargest(maxpackets).plot(kind = 'barh')
		plt.title("Destination IP adresses of packets sent from " + ipadress)
	plt.xlabel("No. of Packets")
	plt.ylabel("IP Adress")
	plt.savefig(savefolder + "/dst_ip.png", bbox_inches='tight')
	plt.clf()
	print("Destination IP Graph Generated")

def sourcePortChart(log_df, ipadress, maxpackets, ip_src_df, ip_dst_df, savefolder):
	plt.figure(figsize=(18,12)) # Set size of all the graphs/maps
	if(ipadress == None): # Generate a bar chart showing the number of packets that were sent from different ports
		log_df['src_port'].value_counts().nlargest(maxpackets).plot(kind = 'barh')
		plt.title("Source ports of packets")
	else: # Generate a bar chart showing the number of packets that were sent from different ports to the specifed ip adress
		ip_dst_df['src_port'].value_counts().nlargest(maxpackets).plot(kind = 'barh')
		plt.title("Source ports of packets sent to " + ipadress)
	plt.xlabel("No. of Packets")
	plt.ylabel("Port")
	plt.savefig(savefolder + "/src_port.png", bbox_inches='tight')
	plt.clf()
	print("Source Port Graph Generated")

def destinationPortChart(log_df, ipadress, maxpackets, ip_src_df, ip_dst_df, savefolder):
	plt.figure(figsize=(18,12)) # Set size of all the graphs/maps
	if(ipadress == None): # Generate a bar chart showing the number of packets that were sent to different ports
		log_df['dst_port'].value_counts().nlargest(maxpackets).plot(kind = 'barh')
		plt.title("Destination ports of packets")
	else: # Generate a bar chart showing the number of packets that were sent to different ports from the specifed ip adress
		ip_src_df['dst_port'].value_counts().nlargest(maxpackets).plot(kind = 'barh')
		plt.title("Destination ports of packets sent from " + ipadress)
	plt.xlabel("No. of Packets")
	plt.ylabel("Port")
	plt.savefig(savefolder + "/dst_port.png", bbox_inches='tight')
	plt.clf()
	
	print("Destination Port Graph Generated")

def sourceIPMap(log_df, ipadress, maxpackets, ip_src_df, ip_dst_df, ip_data, savefolder):
	plt.figure(figsize=(18,12)) # Set size of all the graphs/maps
	m = Basemap(projection='mill',llcrnrlat=-60,urcrnrlat=90,llcrnrlon=-180,urcrnrlon=180,resolution='c')
	m.fillcontinents(color='white',lake_color='black')
	m.drawcountries()
	m.drawmapboundary(fill_color='black')
	plt.xlabel(None)
	plt.ylabel(None)
		
	if(ipadress == None):
		plt.title("Sources of the packets (Orange)")
		ip_src_loc_list = log_df['src_ip'].value_counts().nlargest(maxpackets).index.tolist()
	else:
		plt.title("Sources of the packets (Orange) sent to " + ipadress)# + " (Green)")
		ip_src_loc_list = ip_dst_df['src_ip'].value_counts().nlargest(maxpackets).index.tolist()
			
	for x in ip_src_loc_list:
		if x != "Unknown IP(s)":
			if(ip_data[x].details.get('bogon') != True):
				location = ip_data[x].details.get('loc')
				locationxy = location.split(',')
				m.scatter(float(locationxy[1]), float(locationxy[0]), marker = 'o', color='orange', zorder=5, latlon=True)
	plt.savefig(savefolder + "/src_ip_map.png", bbox_inches='tight', dpi=300)	
	plt.clf()			
	print("Source IP map Generated")

def destinationIPMap(log_df, ipadress, maxpackets, ip_src_df, ip_dst_df, ip_data, savefolder):
	plt.figure(figsize=(18,12)) # Set size of all the graphs/maps
	m = Basemap(projection='mill',llcrnrlat=-60,urcrnrlat=90,llcrnrlon=-180,urcrnrlon=180,resolution='c')
	m.fillcontinents(color='white',lake_color='black')
	m.drawcountries()
	m.drawmapboundary(fill_color='black')
	plt.xlabel(None)
	plt.ylabel(None)
			
	if(ipadress == None):
		plt.title("Destination of the packets (Blue)")
		ip_dst_loc_list = log_df['dst_ip'].value_counts().index.tolist()
	else:
		plt.title("Destination of the packets (Blue) sent to " + ipadress)# + " (Green)")
		ip_dst_loc_list = ip_src_df['dst_ip'].value_counts().index.tolist()
					
	for x in ip_dst_loc_list:
		if x != "Unknown IP(s)":
			if(ip_data[x].details.get('bogon') != True):
				location = ip_data[x].details.get('loc')
				locationxy = location.split(',')
				m.scatter(float(locationxy[1]), float(locationxy[0]), marker = 'o', color='blue', zorder=5, latlon=True)
	plt.savefig(savefolder + "/dst_ip_map.png", bbox_inches='tight', dpi=300)
	plt.clf()
		
	print("Destination IP map Generated")

def documentGeneric(log_df, ipadress, maxpackets, ip_src_df, ip_dst_df, ip_data, savefolder, loadfile, ipinfo_works): # Genrates a document outputting some details about the packet log, TODO cleanup these document, currently messy

	document = Document()

	document.add_heading(loadfile + " Data Analysis", 0) # Set document heading
	document.add_page_break()
	
	srcip_dict = log_df['src_ip'].value_counts().to_dict()
	document.add_heading("Top " + str(maxpackets) + " packet source IP adresses", 1)
	index = 0
	for x in srcip_dict: # Cycle through and print the top x source IPs of packets
		index = index + 1
		document.add_heading(str(index) + ". " + str(x) + " – " + str(srcip_dict[x]) + " packet(s)", 2)
			
		if(ipinfo_works == False or x == "Unknown IP(s)"): # If possable then print their location, else don't
			document.add_paragraph("Location: Unknown")
		elif(ip_data[x].details.get('bogon') == True):
			document.add_paragraph("Location: None/Internal")
		else:
			document.add_paragraph("Location: " + ip_data[x].details.get('city') + ", " + ip_data[x].details.get('region') + ", " + ip_data[x].details.get('country_name') + " " + ip_data[x].details.get('country_flag')['emoji'])
			
		temp_dict = log_df[log_df['src_ip'] == x]['dst_ip'].value_counts().to_dict()
		dest_ip_list = "Top 10 IP adresses that the packets were sent to:\n"
		ipindex = 0
		for y in temp_dict: # Then cycle through and print the top 10 adresses this IP sent packets too
			dest_ip_list += y + " – " + str(temp_dict[y]) + " packet(s).\n"
			ipindex = ipindex + 1
			if(ipindex == 10):
				break
		document.add_paragraph(dest_ip_list) # Add this list to the document
			
		temp_dict = log_df[log_df['src_ip'] == x]['dst_port'].value_counts().to_dict()
		dest_port_list = "Top 10 Ports that the packets were sent to:\n"
		portindex = 0
		for y in temp_dict: # Then cycle through and print the top 10 ports this IP sent packets too
			dest_port_list += str(y) + " – " + str(temp_dict[y]) + " packet(s).\n"
			portindex = portindex + 1
			if(portindex == 10):
				break
		document.add_paragraph(dest_port_list) # Add this list to the document
			
		if (index == maxpackets):
			break
			
	document.add_page_break()
					
	dstip_dict = log_df['dst_ip'].value_counts().to_dict()
	document.add_heading("Top " + str(maxpackets) + " packet destination IP adresses", 1)
	index = 0
	for x in dstip_dict: # Cycle through and print the top x destination IPs of packets
		index = index + 1
		document.add_heading(str(index) + ". " + str(x) + " – " + str(dstip_dict[x]) + " packet(s)", 2)
			
		if(ipinfo_works == False or x == "Unknown IP(s)"): # If possable then print their location, else don't
			document.add_paragraph("Location: Unknown")
		elif(ip_data[x].details.get('bogon') == True):
			document.add_paragraph("Location: None/Internal")
		else:
			document.add_paragraph("Location: " + ip_data[x].details.get('city') + ", " + ip_data[x].details.get('region') + ", " + ip_data[x].details.get('country_name') + " " + ip_data[x].details.get('country_flag')['emoji'])
				
		temp_dict = log_df[log_df['dst_ip'] == x]['src_ip'].value_counts().to_dict()
		dest_ip_list = "Top 10 IP adresses that packets were sent from:\n"
		ipindex = 0
		for y in temp_dict: # Then cycle through and print the top 10 adresses this IP recieved from
			dest_ip_list += y + " – " + str(temp_dict[y]) + " packet(s).\n"
			ipindex = ipindex + 1
			if(ipindex == 10):
				break
			
		document.add_paragraph(dest_ip_list) # Add this list to the document
				
		temp_dict = log_df[log_df['dst_ip'] == x]['src_port'].value_counts().to_dict()
		dest_port_list = "Top 10 Ports that the packets were sent from:\n"
		portindex = 0
		for y in temp_dict: # Then cycle through and print the top 10 ports this IP recieved from
			dest_port_list += str(y) + " – " + str(temp_dict[y]) + " packet(s).\n"
			portindex = portindex + 1
			if(portindex == 10):
				break
		document.add_paragraph(dest_port_list) # Add this list to the document
				
		if (index == maxpackets):
			break
			
	document.add_page_break()
	
	document.save(savefolder + "/Analysis Document.docx") # Save document
	
	print("Document Generated")
		
def documentSpecific(log_df, ipadress, maxpackets, ip_src_df, ip_dst_df, ip_data, savefolder, loadfile, ipinfo_works): # Genrates a document outputting some details about a specific IP adress in the packet log, TODO clean up and possibly combine with documentGeneric to reduce code duplication
	
	document = Document()
	
	document.add_heading(loadfile + " Data Analysis for " + ipadress, 0) # Set document heading
		
	if(ipinfo_works == False or ipadress == "Unknown IP(s)"): # If possable then print the specifed IP adress, coordiantes, ISP, host, else dont
			document.add_paragraph("Location: Unknown")
	elif(ip_data[ipadress].details.get('bogon') != True):
		document.add_paragraph("Location: " + ip_data[ipadress].details.get('city') + ", " + ip_data[ipadress].details.get('region') + ", " + ip_data[ipadress].details.get('country_name') + " " + ip_data[ipadress].details.get('country_flag')['emoji'])
		document.add_paragraph("Coordinates: " + ip_data[ipadress].details.get('loc'))
		document.add_paragraph("ISP: " + ip_data[ipadress].details.get('org'))
		if (ip_data[ipadress].details.get('hostname') != None):
			document.add_paragraph("Host: " + ip_data[ipadress].details.get('hostname'))
	else:
		document.add_paragraph("This IP adress is bogon and dose not have any assigned data, it is likely used by a local network.")
	
	document.add_page_break()
		
	temp_dict = ip_src_df['dst_ip'].value_counts().to_dict()
	document.add_heading("IP adresses that " + ipadress + " sent packets to:\n", 1)
	dest_ip_list = ""
	for x in temp_dict: # Print every IP adress that recieved packets from the specifed IP adress
		dest_ip_list += x + " – " + str(temp_dict[x]) + " packet(s).\n"	
	document.add_paragraph(dest_ip_list) # Add this list to the document
	document.add_page_break()
		
	temp_dict = ip_dst_df['src_ip'].value_counts().to_dict()
	document.add_heading("IP adresses that " + ipadress + " has recieved packets from:\n", 1)
	src_ip_list = ""
	for x in temp_dict: # Print every IP adress that sent packets to the specifed IP adress
		src_ip_list += x + " – " + str(temp_dict[x]) + " packet(s).\n"	
	document.add_paragraph(dest_ip_list) # Add this list to the document
	document.add_page_break()
		
	temp_dict = ip_src_df['dst_port'].value_counts().to_dict()
	document.add_heading("Ports that " + ipadress + " sent packets to:\n", 1)
	dest_ip_list = ""
	for x in temp_dict: # Print every port that recieved packets from the specifed IP adress
		dest_ip_list += str(x) + " – " + str(temp_dict[x]) + " packet(s).\n"	
	document.add_paragraph(dest_ip_list) # Add this list to the document
	document.add_page_break()
		
	temp_dict = ip_dst_df['src_port'].value_counts().to_dict()
	document.add_heading("Ports that " + ipadress + " has recieved packets from:\n", 1)
	src_ip_list = ""
	for x in temp_dict: # Print every port that sent packets to the specifed IP adress
		src_ip_list += str(x) + " – " + str(temp_dict[x]) + " packet(s).\n"	
	document.add_paragraph(dest_ip_list) # Add this list to the document
	document.add_page_break()
		
	document.save(savefolder + "/Analysis Document.docx") # Save the Document
	
	print("Document Generated")